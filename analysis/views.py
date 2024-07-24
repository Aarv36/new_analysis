from django.shortcuts import render, redirect, get_object_or_404
from django.conf import settings
from django.contrib import messages
from .forms import UploadFileForm
from .models import UploadedFile
from .utils import analyze_pdf, extract_text_from_pdf, extract_text_from_word, find_details_in_text, save_text_to_word
import os
import shutil
import uuid
from django.http import HttpResponse, FileResponse, Http404
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from io import BytesIO
import zipfile
import logging
from django.http import FileResponse
import io
import cv2
from docx import Document
from docx.shared import RGBColor

# Configure logging
logger = logging.getLogger(__name__)

def upload_file(request):
    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                uploaded_file = request.FILES['file']
                logger.debug(f"Uploaded file: {uploaded_file.name}")
                uploaded_file_instance = UploadedFile(file=uploaded_file)
                uploaded_file_instance.save()

                file_path = uploaded_file_instance.file.path
                logger.info(f"Uploaded file path: {file_path}")

                if file_path.endswith('.pdf'):
                    extracted_text = extract_text_from_pdf(file_path)
                elif file_path.endswith('.docx'):
                    extracted_text = extract_text_from_word(file_path)
                elif file_path.endswith('.jpg') or file_path.endswith('.jpeg'):
                    # Directly analyze the image for QR and Aadhar details
                    pass
                else:
                    messages.error(request, "Unsupported file type")
                    return redirect('upload_file')

                extracted_details, modified_text = find_details_in_text(extracted_text)

                # Ensure extracted_details is a dictionary
                if isinstance(extracted_details, dict):
                    logger.info(f"Extracted details: {extracted_details}")
                else:
                    raise ValueError("Expected extracted_details to be a dictionary")

                request.session['modified_text'] = modified_text
                request.session['extracted_details'] = extracted_details
                # Save modified text to a Word file
                word_file_path = os.path.join(settings.MEDIA_ROOT, 'modified_text',
                                              f'modified_text_{uuid.uuid4().hex}.docx')
                word_file_dir = os.path.dirname(word_file_path)
                if not os.path.exists(word_file_dir):
                    os.makedirs(word_file_dir)
                save_text_to_word(modified_text, word_file_path)
                request.session['word_file_path'] = word_file_path

                logger.debug("Redirecting to analyze_file view")
                return redirect('analyze_file', file_id=uploaded_file_instance.id)
            except Exception as e:
                logger.error(f"Error in upload_file: {e}")
                messages.error(request, f"Error: {e}")
        else:
            messages.error(request, "Invalid form submission")
    else:
        logger.debug("GET request received")
        form = UploadFileForm()
    return render(request, 'analysis/upload.html', {'form': form})


def download_highlighted_images(request, file_id, format):
    uploaded_file = get_object_or_404(UploadedFile, id=file_id)
    output_folder = os.path.join(settings.MEDIA_ROOT, 'temp', str(file_id))
    highlighted_images = [os.path.join(output_folder, f) for f in os.listdir(output_folder) if f.endswith('.png')]

    if format == 'pdf':
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="highlighted_images_{file_id}.pdf"'

        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)

        for img_path in highlighted_images:
            p.drawImage(img_path, 0, 0, width=letter[0], height=letter[1])
            p.showPage()

        p.save()
        pdf = buffer.getvalue()
        buffer.close()
        response.write(pdf)
        return response

    elif format == 'zip':
        response = HttpResponse(content_type='application/zip')
        response['Content-Disposition'] = f'attachment; filename="highlighted_images_{file_id}.zip"'

        buffer = BytesIO()
        with zipfile.ZipFile(buffer, 'w') as zip_file:
            for img_path in highlighted_images:
                zip_file.write(img_path, os.path.basename(img_path))
        buffer.seek(0)
        response.write(buffer.read())
        return response

    return HttpResponse(status=400)


def download_word_file(request, file_id):
    modified_text = request.session.get('modified_text', "No text found")
    if not modified_text:
        return HttpResponse(status=404)

    # Create a new Document
    doc = Document()
    doc.add_heading('Modified Text Content', level=1)
    #doc.add_paragraph(modified_text)

    # Split text into parts to identify and style `xxxxx`
    parts = modified_text.split('xxxxx')
    for i, part in enumerate(parts):
        run = doc.add_paragraph().add_run(part)
        if i < len(parts) - 1:
            xxxxx_run = doc.add_paragraph().add_run('xxxxx')
            xxxxx_run.bold = True
            xxxxx_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color

    # Save the document to a BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Create a FileResponse
    response = FileResponse(buffer, as_attachment=True, filename=f'modified_text_{file_id}.docx')
    return response

def analyze_file(request, file_id):
    logger.debug(f"Entered analyze_file view with file_id: {file_id}")
    try:
        uploaded_file = get_object_or_404(UploadedFile, id=file_id)
        file_path = os.path.join(settings.MEDIA_ROOT, uploaded_file.file.name)
        #logger.info(f"Analyzing file path: {file_path}")

        if file_path.endswith('.pdf'):
            extracted_text = extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            extracted_text = extract_text_from_word(file_path)
        else:
            return HttpResponse('Unsupported file type.', status=400)


        output_folder = os.path.join(settings.MEDIA_ROOT, 'temp', str(file_id))
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        message, highlighted_pages, aadhar_details, face_details, page_statuses = analyze_pdf(file_path, output_folder)
        highlighted_images_urls = []

        for image_path in highlighted_pages:
            base_name = os.path.basename(image_path)
            unique_name = f"{uuid.uuid4().hex}_{base_name}"
            destination = os.path.join(output_folder, unique_name)
            shutil.move(image_path, destination)
            highlighted_images_urls.append(os.path.join(settings.MEDIA_URL, 'temp', str(file_id), unique_name))


        extracted_details = request.session.pop('extracted_details', {})
        extracted_text = request.session.pop('extracted_text', "")
        modified_text = request.session.get('modified_text', "No text found")
        logger.debug(f"Anonymized text: {modified_text}")

        # Combine details by page
        combined_details = {}
        for page_num, details in aadhar_details:
            if page_num not in combined_details:
                combined_details[page_num] = {'aadhar': [], 'faces': [], 'status': {}}
            combined_details[page_num]['aadhar'].extend(details)

        for page_num, faces in face_details:
            if page_num not in combined_details:
                combined_details[page_num] = {'aadhar': [], 'faces': [], 'status': {}}
            combined_details[page_num]['faces'].extend(faces)

        # Add statuses to combined details
        for page_num, status in page_statuses.items():
            if page_num not in combined_details:
                combined_details[page_num] = {'aadhar': [], 'faces': [], 'status': {}}
            combined_details[page_num]['status'] = status

        logger.debug(f"Rendering result.html with context: {locals()}")
        return render(request, 'analysis/result.html', {
            'message': message,
            'highlighted_images': highlighted_images_urls,
            'combined_details': combined_details,
            'extracted_details': extracted_details,
            'modified_text': modified_text,
            'file_id': file_id,

        })
    except Exception as e:
        logger.error(f"Error in analyze_file: {e}")
        messages.error(request, f"Error: {e}")
        return redirect('upload_file')
