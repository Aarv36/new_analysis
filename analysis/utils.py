import re
import easyocr
import cv2
import numpy as np
import os
import time
import copy
import PyPDF2
from docx import Document

reader = easyocr.Reader(['en'])

current_dir = os.path.dirname(os.path.abspath(__file__))
yolo_cfg_path = os.path.join(current_dir, "yolov3.cfg")
yolo_weights_path = os.path.join(current_dir, "yolov3last2.weights")

net = cv2.dnn.readNet(yolo_weights_path, yolo_cfg_path)
classes = ["name", 'dob', 'gender', 'aadhar_no']
layer_names = net.getLayerNames()
output_layers = [layer_names[i - 1] for i in net.getUnconnectedOutLayers()]
colors = np.random.uniform(0, 255, size=(len(classes), 3))

# Add the path to the Haar Cascade XML file for face detection
face_cascade_path = os.path.join(current_dir, "haarcascade_frontalface_default.xml")
face_cascade = cv2.CascadeClassifier(face_cascade_path)

def validate_aadhar_details(details):
    validated_details = []
    for detail_type, texts, bbox in details:
        if detail_type == 'aadhar_no':
            validated_texts = [text for text in texts if re.match(r'^\d{4}\s\d{4}\s\d{4}$', text)]
        elif detail_type == 'dob':
            validated_texts = [text for text in texts if re.match(r'^\d{2}/\d{2}/\d{4}$', text)]
        else:
            validated_texts = texts

        if validated_texts:
            validated_details.append((detail_type, validated_texts, bbox))

    return validated_details


def convert_pdf_to_images(pdf_path, dpi=300):
    import fitz  # PyMuPDF
    pdf_document = fitz.open(pdf_path)
    image_list = []

    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        pix = page.get_pixmap(dpi=dpi)
        image_data = pix.tobytes("png")
        image = cv2.imdecode(np.frombuffer(image_data, np.uint8), cv2.IMREAD_COLOR)
        image_list.append(image)

    return image_list


def detect_qr_in_image(image):
    qr_detector = cv2.QRCodeDetector()
    data, points, _ = qr_detector.detectAndDecode(image)
    return data, points

def detect_faces(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    faces = face_cascade.detectMultiScale(gray, scaleFactor=1.1, minNeighbors=5, minSize=(30, 30))
    return faces

def draw_faces(image, faces):
    for (x, y, w, h) in faces:
        cv2.rectangle(image, (x, y), (x + w, y + h), (255, 0, 0), 2)
        # Add filled rectangle
        cv2.rectangle(image, (x, y), (x + w, y + h), (255, 0, 0), -1)  # -1 fills the rectangle
    return image

def detect_aadhar_details(image):
    blob = cv2.dnn.blobFromImage(image, 0.00392, (416, 416), (0, 0, 0), True, crop=False)
    net.setInput(blob)
    outs = net.forward(output_layers)

    class_ids = []
    confidences = []
    boxes = []
    details = []

    for out in outs:
        for detection in out:
            scores = detection[5:]
            class_id = np.argmax(scores)
            confidence = scores[class_id]
            if confidence > 0.25:
                center_x = int(detection[0] * image.shape[1])
                center_y = int(detection[1] * image.shape[0])
                w = int(detection[2] * image.shape[1])
                h = int(detection[3] * image.shape[0])
                x = int(center_x - w / 2)
                y = int(center_y - h / 2)
                boxes.append([x, y, w, h])
                confidences.append(float(confidence))
                class_ids.append(class_id)

    indexes = cv2.dnn.NMSBoxes(boxes, confidences, 0.5, 0.4)
    for i in range(len(boxes)):
        if i in indexes:
            x, y, w, h = boxes[i]
            crop = image[y:y + h, x:x + w]
            result = reader.readtext(crop, detail=0)
            details.append((classes[class_ids[i]], result, (x, y, w, h)))

    validated_details = validate_aadhar_details(details)

    return validated_details


def extract_text_from_pdf(pdf_path):
    pdf_reader = PyPDF2.PdfReader(pdf_path)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def extract_text_from_word(docx_path):
    doc = Document(docx_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


def find_details_in_text(text):
    details = {"names": [],
               "aadhar_numbers": [],
               "dobs": [],
               "pan_numbers": []
               }

    # Regular expressions for matching patterns
    name_pattern = re.compile(r'\b[A-Z][a-z]+\s[A-Z][a-z]+\b')  # Simplified name pattern
    aadhar_pattern = re.compile(r'\b\d{4}[\s.,/(){}\[\]\\|!*&^_+#@]?\d{4}[\s.,/(){}\[\]\\|!*&^_+#@]?\d{4}\b')
    dob_pattern = re.compile(r'\b(?:\d{4}[-/]\d{2}[-/]\d{2}|\d{2}[-/]\d{2}[-/]\d{4}|\d{2}[-/]\d{4}[-/]\d{2})\b')
    pan_pattern = re.compile(r'\b[A-Za-z]{5}[0-9]{4}[A-Za-z]{1}\b')

    details["names"] = name_pattern.findall(text)
    details["aadhar_numbers"] = aadhar_pattern.findall(text)
    details["dobs"] = dob_pattern.findall(text)
    details["pan_numbers"] = pan_pattern.findall(text)

    # Replace detected details with "XXXXX"
    text = name_pattern.sub("XXXXX", text)
    text = aadhar_pattern.sub("XXXXX", text)
    text = dob_pattern.sub("XXXXX", text)
    text = pan_pattern.sub("XXXXX", text)

    return details, text

def save_text_to_word(text, file_path):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(file_path)


def draw_filled_boxes(image, details):
    for detail in details:
        print(detail)  # Debug: print details
        if len(detail) == 3:
            detail_type, texts, (x, y, w, h) = detail
            color = colors.get(detail_type, (0, 0, 255))  # Default to blue if type not found
            overlay = image.copy()
            alpha = 0.4  # Transparency factor

            # Draw the filled rectangle with transparency
            cv2.rectangle(overlay, (x, y), (x + w, y + h), color, -1)
            # Add the rectangle to the image with transparency
            image = cv2.addWeighted(overlay, alpha, image, 1 - alpha, 0)

            # Draw the detail type and text on the filled box
            cv2.putText(image, detail_type, (x, y - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.9, (255, 255, 255), 2)
            cv2.putText(image, ', '.join(texts), (x, y + h + 20), cv2.FONT_HERSHEY_SIMPLEX, 0.6, (255, 255, 255), 2)
    return image


def draw_tags(image, qrcode_result, aadhar_results, elapsed_time):
    if qrcode_result[0]:
        text = qrcode_result[0]
        points = qrcode_result[1]
        if points is not None:
            points = points[0]

            # Convert points to integer tuples
            int_points = [tuple(map(int, pt)) for pt in points]

            # Draw a filled polygon for half of the QR code
            # Assuming "half" means the upper half of the QR code bounding box
            top_points = int_points[:2] + [
                (int_points[1][0], int_points[1][1] + (int_points[2][1] - int_points[1][1]) // 2),
                (int_points[0][0], int_points[0][1] + (int_points[2][1] - int_points[1][1]) // 2)]

            # Draw the filled polygon
            cv2.fillPoly(image, [np.array(top_points)], (255, 0, 0))  # Fill color is blue

            # Optionally, you can fill the entire QR code box by commenting the above lines and uncommenting the line below
            # cv2.fillPoly(image, [np.array(int_points)], (255, 0, 0))  # Full fill color is blue

            # Draw the QR code box lines over the filled polygon
            for i in range(len(int_points)):
                pt1 = int_points[i]
                pt2 = int_points[(i + 1) % len(int_points)]
                cv2.line(image, pt1, pt2, (0, 255, 0), 3)

            # Draw the QR code text
            cv2.putText(image, text, (int(int_points[0][0]), int(int_points[0][1]) - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.9,
                        (0, 255, 0), 2)


    for detail in aadhar_results:
        print(f"Processing aadhar detail: {detail}")
        try:
            if len(detail) == 3:
                detail_type, texts, (x, y, w, h) = detail
                color = colors[classes.index(detail_type)]
                # Draw filled rectangle first
                cv2.rectangle(image, (x, y), (x + w, y + h), color, -1)  # -1 fills the rectangle
                # Draw the border rectangle
                cv2.rectangle(image, (x, y), (x + w, y + h), color, 2)
                cv2.putText(image, detail_type, (x, y - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.9, color, 2)
                cv2.putText(image, ', '.join(texts), (x, y + h + 20), cv2.FONT_HERSHEY_SIMPLEX, 0.6, color, 2)
            else:
                print(f"Unexpected detail format in aadhar_results: {detail}")
        except ValueError as e:
            print(f"Error unpacking aadhar detail: {detail}, Error: {e}")

    cv2.putText(image, f"Elapsed Time: {elapsed_time * 1000:.1f} ms", (10, 30), cv2.FONT_HERSHEY_SIMPLEX, 0.8,
                (0, 255, 0), 2, cv2.LINE_AA)
    return image

def resize_image_to_fit_screen(image, screen_width, screen_height):
    height, width = image.shape[:2]
    scaling_factor = min(screen_width / width, screen_height / height)
    new_size = (int(width * scaling_factor), int(height * scaling_factor))
    resized_image = cv2.resize(image, new_size, interpolation=cv2.INTER_AREA)
    return resized_image


def analyze_pdf(pdf_path, output_folder, screen_width=1280, screen_height=720):
    print("Starting analyze_pdf function")
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    images = convert_pdf_to_images(pdf_path)
    qr_detected = False
    aadhar_details = []
    highlighted_pages = []
    face_details = []

    page_statuses = {}


    for page_num, image in enumerate(images, start=1):
        page_statuses[page_num] = {
            'qr': False,
            'aadhar': False,
            'faces': False,

        }
        start_time = time.time()
        qrcode_result = detect_qr_in_image(image)
        aadhar_result = detect_aadhar_details(image)
        faces = detect_faces(image)
        elapsed_time = time.time() - start_time


        print(f"Processing page {page_num}, elapsed time: {elapsed_time:.2f} seconds")
        # Print the function call details
        print("Calling draw_tags with the following parameters:")
        print(
            f"Image: {type(image)}, QRCode Result: {qrcode_result}, Aadhar Result: {aadhar_result}, Elapsed Time: {elapsed_time}")

        try:
            debug_image = draw_tags(copy.deepcopy(image), qrcode_result, aadhar_result, elapsed_time)
        except Exception as e:
            print(f"Error in draw_tags: {e}")
            raise

        debug_image_with_faces = draw_faces(debug_image, faces)
        resized_image = resize_image_to_fit_screen(debug_image, screen_width, screen_height)
        highlighted_page_path = os.path.join(output_folder, f"page_{page_num}.png")
        cv2.imwrite(highlighted_page_path, resized_image)
        highlighted_pages.append(highlighted_page_path)

        if qrcode_result[0]:
            qr_detected = True
            page_statuses[page_num]['qr'] = True

        if aadhar_result:
            aadhar_details.append((page_num, aadhar_result))
            page_statuses[page_num]['aadhar'] = True

        if len(faces) > 0:
            face_details.append((page_num, faces))
            page_statuses[page_num]['faces'] = True


    result_message = []
    if qr_detected:
        result_message.append("QR code detected in the PDF.")
    else:
        result_message.append("No QR code found in the PDF.")

    if aadhar_details:
        result_message.append("Aadhar details detected in the PDF.")
    else:
        result_message.append("No Aadhar details found in the PDF.")

    if face_details:
        result_message.append("Faces detected in the PDF.")
    else:
        result_message.append("No faces found in the PDF.")

    return " ".join(result_message), highlighted_pages, aadhar_details, face_details, page_statuses
